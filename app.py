# ❌ REMOVE THIS PART (for Streamlit Cloud):
# from dotenv import load_dotenv
# load_dotenv()
# GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# ✅ USE THIS INSTEAD:
import streamlit as st
import os
import google.generativeai as genai
from docx import Document
from docx.shared import RGBColor, Pt
from io import BytesIO

# --- PAGE CONFIGURATION ---
st.set_page_config(
    page_title="AI Content Studio",
    page_icon="✨",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- CUSTOM CSS (same as before) ---
st.markdown("""
<style>
    .stApp {
        background-color: #0f172a;
        color: #f8fafc;
    }
    
    [data-testid="stSidebar"] {
        background-color: #1e293b;
        border-right: 1px solid #334155;
    }
    
    h1, h2, h3 {
        color: white !important;
        font-family: 'Segoe UI', sans-serif;
    }

    .stTextInput > div > div > input,
    .stSelectbox > div > div > select {
        background-color: #0f172a;
        color: white;
        border: 1px solid #475569;
    }

    .stButton > button {
        background: linear-gradient(90deg, #6366f1, #ec4899);
        color: white;
        border-radius: 8px;
        border: none;
        font-weight: bold;
        padding: 0.5rem 1rem;
        transition: all 0.3s;
        width: 100%;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 5px 15px rgba(99, 102, 241, 0.4);
    }

    .stDownloadButton > button {
        background-color: #10b981;
        color: white;
        border-radius: 8px;
        font-weight: bold;
        width: 100%;
        border: none;
    }
    
    .stTextArea > div > div > textarea {
        background-color: #1e293b;
        color: #e2e8f0;
        border: 1px solid #334155;
    }
</style>
""", unsafe_allow_html=True)

# --- SECURE API KEY LOADING ---
# This works for both local (with .env) AND Streamlit Cloud (with secrets)
def get_api_key():
    """Get API key from Streamlit secrets or environment variables"""
    try:
        # Try Streamlit Cloud secrets first
        return st.secrets["GOOGLE_API_KEY"]
    except KeyError:
        # Fall back to environment variable (for local development)
        api_key = os.getenv("GOOGLE_API_KEY")
        if api_key:
            return api_key
        else:
            return None

GOOGLE_API_KEY = get_api_key()

if not GOOGLE_API_KEY:
    st.error("🚨 **API Key Missing!**\n\n"
             "**For Streamlit Cloud:** Add `GOOGLE_API_KEY` in Secrets.\n"
             "**For Local:** Create `.env` file with `GOOGLE_API_KEY=your_key`")
    st.stop()

# --- AI FUNCTIONS ---
@st.cache_resource
def load_model():
    genai.configure(api_key=GOOGLE_API_KEY)
    return genai.GenerativeModel('gemini-2.5-flash')

def create_prompt(topic, category, platform, tone):
    instructions = {
        "YouTube": "Write a script with visual cues in [brackets].",
        "LinkedIn": "Professional tone, under 1500 chars, include hashtags.",
        "Instagram / TikTok": "Short, punchy, emoji-rich, under 60 seconds.",
        "Personal Blog": "Write an engaging, SEO-optimized article with H2 headers.",
        "default": "Write a comprehensive article."
    }
    
    inst = instructions.get(platform, instructions["default"])
    return f"Write a {category} about '{topic}'. Tone: {tone}. Platform: {platform}. {inst}"

def generate_content(prompt):
    try:
        model = load_model()
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error generating content: {e}"

def create_docx(content, filename):
    doc = Document()
    
    title_text = filename.replace(".docx", "")
    heading = doc.add_heading(title_text, level=1)
    heading.style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    
    doc.add_paragraph(content)
    
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

# --- DASHBOARD UI ---
st.title("✨ AI Content Studio")
st.markdown("---")

with st.sidebar:
    st.header("⚙️ Configuration")
    
    topic = st.text_input("Content Topic", placeholder="e.g., The Future of AI")
    category = st.selectbox("Category", ["Blog Post", "Video Script", "News Article", "Social Media Caption", "Technical Article"])
    platform = st.selectbox("Platform", ["YouTube", "LinkedIn", "Instagram / TikTok", "Personal Blog", "Medium"])
    tone = st.selectbox("Tone of Voice", ["Professional", "Casual & Witty", "Educational", "Persuasive", "Urgent"])
    
    st.markdown("---")
    generate_btn = st.button("✨ Generate Content", use_container_width=True)

col1, col2 = st.columns([2, 1])

with col1:
    st.subheader("📄 Content Preview")
    
    if "content" not in st.session_state:
        st.session_state.content = ""
    
    if generate_btn:
        if not topic:
            st.warning("Please enter a topic first.")
        else:
            with st.spinner("✍️ AI is writing your content..."):
                prompt = create_prompt(topic, category, platform, tone)
                st.session_state.content = generate_content(prompt)
    
    if st.session_state.content:
        st.text_area("Generated Output", st.session_state.content, height=400)
    else:
        st.info("👈 Fill out the sidebar and click Generate to start.")

with col2:
    st.subheader("📥 Export Options")
    
    if st.session_state.content:
        st.success("Content generated successfully!")
        
        safe_filename = f"{topic.replace(' ', '_')}.docx"
        doc_file = create_docx(st.session_state.content, safe_filename)
        
        st.download_button(
            label="Download Word Document (.docx)",
            data=doc_file,
            file_name=safe_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        st.markdown("---")
        st.caption(f"Model: gemini-2.5-flash\nPlatform: {platform}\nTone: {tone}")
    else:
        st.caption("Waiting for content generation...")